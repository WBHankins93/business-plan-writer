"""Build the fictional public sample plan as a polished DOCX.

The source is deliberately plain Markdown so the fictional content stays easy to
audit. This builder applies the neighborhood_business_proposal document preset
and proposal_centerpiece cover pattern used for the public sample download.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "web" / "public" / "samples" / "bywater-grounds-sample-plan.md"
OUTPUT = ROOT / "web" / "public" / "samples" / "bywater-grounds-sample-plan.docx"
DOC_SKILL = Path(
    "/Users/benhankins/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.715.12143/skills/documents"
)
sys.path.insert(0, str(DOC_SKILL))

from scripts.table_geometry import apply_table_geometry  # noqa: E402


NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x5D, 0x69, 0x70)
LIGHT_FILL = "F4F6F9"
BODY_FONT = "Calibri"


def set_run_font(run, *, size: float, color=RGBColor(0, 0, 0), bold=False, italic=False) -> None:
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=8, color=GRAY)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    note_style = doc.styles.add_style("Sample Notice", WD_STYLE_TYPE.PARAGRAPH)
    note_style.base_style = normal
    note_style.font.name = BODY_FONT
    note_style.font.size = Pt(10)
    note_style.font.color.rgb = GRAY
    note_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_style.paragraph_format.space_before = Pt(12)
    note_style.paragraph_format.space_after = Pt(8)
    note_style.paragraph_format.line_spacing = 1.25

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_p.add_run("BYWATER GROUNDS  |  FICTIONAL SAMPLE")
    set_run_font(header_run, size=8, color=GRAY, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer_p.add_run("Business Plan Writer  |  ")
    set_run_font(label, size=8, color=GRAY)
    add_page_field(footer_p)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(104)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    set_run_font(kicker.add_run("FICTIONAL SAMPLE  /  PRIVATE BETA"), size=9, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(7)
    set_run_font(title.add_run("Bywater Grounds\nCoffee House"), size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_run_font(subtitle.add_run("Representative lender-ready business plan"), size=14, color=GRAY)

    request = doc.add_paragraph()
    request.alignment = WD_ALIGN_PARAGRAPH.CENTER
    request.paragraph_format.space_after = Pt(4)
    set_run_font(request.add_run("SBA 7(a) funding request  |  $285,000"), size=11, color=DARK_BLUE, bold=True)

    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared.paragraph_format.space_after = Pt(68)
    set_run_font(prepared.add_run("Prepared July 2026"), size=10, color=GRAY)

    notice = doc.add_paragraph(style="Sample Notice")
    notice.add_run(
        "This made-up business and every figure in this document exist only to demonstrate "
        "the expected structure and formatting. No external research or customer data was used."
    )
    doc.add_page_break()


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if bold else part
        run = paragraph.add_run(clean)
        set_run_font(run, size=11, bold=bold)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for row_index, row_data in enumerate(rows):
        for column_index, value in enumerate(row_data[:2]):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(value.replace("**", ""))
            set_run_font(run, size=9.5, bold=(row_index == 0 or "**" in value))
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
    apply_table_geometry(
        table,
        [6500, 2860],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_source(doc: Document, source: str) -> None:
    lines = source.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            index += 1
            continue
        if line.startswith("|") and line.endswith("|"):
            table_rows: list[list[str]] = []
            while index < len(lines):
                table_line = lines[index].strip()
                if not (table_line.startswith("|") and table_line.endswith("|")):
                    break
                values = [value.strip() for value in table_line.strip("|").split("|")]
                if not all(re.fullmatch(r"[-: ]+", value) for value in values):
                    table_rows.append(values)
                index += 1
            add_table(doc, table_rows)
            continue
        paragraph = doc.add_paragraph()
        add_inline_markdown(paragraph, line)
        index += 1


def build() -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    parse_source(doc, SOURCE.read_text(encoding="utf-8"))
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
