"""
output/docx_builder.py
Converts the markdown business plan into a professionally styled Word document.

Handles:
  - ## Section headers → Heading 1 style
  - ### Subsection headers → Heading 2 style
  - **bold** inline formatting
  - | table | rows | → Word tables
  - [WRITER_NOTE: ...] → highlighted notes for human review
  - Regular paragraphs → Normal style with appropriate spacing
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Style constants ────────────────────────────────────────────────────────────
_FONT_BODY = "Calibri"
_FONT_HEADING = "Calibri"
_FONT_SIZE_BODY = 11
_FONT_SIZE_H1 = 16
_FONT_SIZE_H2 = 13
_COLOR_H1 = RGBColor(0x1F, 0x49, 0x7D)      # dark navy
_COLOR_H2 = RGBColor(0x2E, 0x74, 0xB5)      # medium blue
_COLOR_NOTE = RGBColor(0xC0, 0x50, 0x00)    # amber/orange for writer notes
_COLOR_TABLE_HEADER = RGBColor(0x1F, 0x49, 0x7D)

# Regex for parsing
_RE_H1 = re.compile(r"^##\s+(.+)$")
_RE_H2 = re.compile(r"^###\s+(.+)$")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_TABLE_SEP = re.compile(r"^\|[-| :]+\|$")
_RE_WRITER_NOTE = re.compile(r"\[WRITER_NOTE[:\s]([^\]]*)\]", re.IGNORECASE)
_RE_BOLD = re.compile(r"\*\*(.+?)\*\*")


def build_docx(
    business_plan_markdown: str,
    business_name: str,
    output_path: str | Path,
    intake_date: str | None = None,
) -> Path:
    """
    Convert a markdown business plan into a styled .docx file.

    Args:
        business_plan_markdown: Full markdown text from Agent 4.
        business_name: Used in the document header/title.
        output_path: Where to save the .docx file.
        intake_date: Optional date string for the document footer.

    Returns:
        Path to the saved .docx file.
    """
    doc = Document()
    _configure_page(doc)
    _add_title_block(doc, business_name, intake_date)
    _parse_and_write(doc, business_plan_markdown)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ── Document setup ─────────────────────────────────────────────────────────────

def _configure_page(doc: Document) -> None:
    """Set margins and default paragraph font."""
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Default paragraph style
    style = doc.styles["Normal"]
    font = style.font
    font.name = _FONT_BODY
    font.size = Pt(_FONT_SIZE_BODY)


def _add_title_block(doc: Document, business_name: str, intake_date: str | None) -> None:
    """Add the document title and subtitle."""
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("BUSINESS PLAN")
    title_run.font.name = _FONT_HEADING
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = _COLOR_H1
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run(business_name)
    sub_run.font.name = _FONT_HEADING
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = _COLOR_H2
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_str = intake_date or date.today().strftime("%B %d, %Y")
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Prepared: {date_str}")
    date_run.font.name = _FONT_BODY
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer


# ── Markdown parser / writer ───────────────────────────────────────────────────

def _parse_and_write(doc: Document, markdown: str) -> None:
    """Parse the markdown and write structured content to the document."""
    lines = markdown.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # H1 header (##)
        m = _RE_H1.match(line)
        if m:
            _add_heading(doc, m.group(1), level=1)
            i += 1
            continue

        # H2 header (###)
        m = _RE_H2.match(line)
        if m:
            _add_heading(doc, m.group(1), level=2)
            i += 1
            continue

        # Table — collect all table rows
        if _RE_TABLE_ROW.match(line):
            table_lines = []
            while i < len(lines) and _RE_TABLE_ROW.match(lines[i]):
                if not _RE_TABLE_SEP.match(lines[i]):
                    table_lines.append(lines[i])
                i += 1
            if table_lines:
                _add_table(doc, table_lines)
            continue

        # Empty line — skip or add spacing
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        _add_paragraph(doc, line.strip())
        i += 1


def _add_heading(doc: Document, text: str, level: int) -> None:
    """Add a styled heading paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = _FONT_HEADING
    run.font.bold = True

    if level == 1:
        run.font.size = Pt(_FONT_SIZE_H1)
        run.font.color.rgb = _COLOR_H1
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(8)
    else:
        run.font.size = Pt(_FONT_SIZE_H2)
        run.font.color.rgb = _COLOR_H2
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)


def _add_paragraph(doc: Document, text: str) -> None:
    """Add a body paragraph, handling bold and WRITER_NOTE tags."""
    # Check for WRITER_NOTE
    note_match = _RE_WRITER_NOTE.search(text)

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)

    if note_match:
        # Render the note in amber color
        note_text = f"[Note for review: {note_match.group(1).strip()}]"
        before = text[: note_match.start()].strip()
        after = text[note_match.end() :].strip()

        if before:
            _add_inline_runs(para, before)
        note_run = para.add_run(f" {note_text} " if before or after else note_text)
        note_run.font.color.rgb = _COLOR_NOTE
        note_run.font.italic = True
        note_run.font.name = _FONT_BODY
        note_run.font.size = Pt(_FONT_SIZE_BODY)
        if after:
            _add_inline_runs(para, after)
    else:
        _add_inline_runs(para, text)


def _add_inline_runs(para, text: str) -> None:
    """Split text on **bold** markers and add runs with appropriate formatting."""
    parts = _RE_BOLD.split(text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.name = _FONT_BODY
        run.font.size = Pt(_FONT_SIZE_BODY)
        # Odd-indexed parts are bold (inside ** **)
        run.bold = (idx % 2 == 1)


def _add_table(doc: Document, table_lines: list[str]) -> None:
    """Parse and add a markdown table to the document."""
    rows = []
    for line in table_lines:
        # Strip leading/trailing | and split
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_cells in enumerate(rows):
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx >= col_count:
                break
            cell = table.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.name = _FONT_BODY
            run.font.size = Pt(10)

            # Style the header row
            if r_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Set cell background color
                _set_cell_background(cell, "1F497D")

    doc.add_paragraph()  # spacing after table


def _set_cell_background(cell, hex_color: str) -> None:
    """Set a table cell's background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in(
        "w:shd",
    )
    if shd is None:
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
